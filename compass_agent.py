# Importing necessary libraries for the system
import os
__import__('pysqlite3')
import sys
sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')
import streamlit as st
import pandas as pd
import docx
from docx import Document
from typing import Dict, List, Optional
import chromadb
from chromadb.utils import embedding_functions
from openai import OpenAI
import requests
import logging
import json
from datetime import datetime
from io import BytesIO

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Global variables for database collections
university_collection = None
living_expenses_collection = None
employment_collection = None
chroma_client = None
embedding_function = None

# Location to city mapping for weather
LOCATION_TO_CITY = {
    "Northeast": "New York",
    "Southeast": "Miami",
    "Midwest": "Chicago",
    "Southwest": "Houston",
    "West Coast": "Los Angeles",
}

# New global variables for enhanced features
ALLOWED_TOPICS = [
    "university",
    "college",
    "study",
    "education",
    "program",
    "course",
    "major",
    "weather",
    "climate",
    "temperature",
    "scholarship",
    "financial aid",
    "funding",
    "loan",
    "job",
    "career",
    "employment",
    "industry",
    "economy",
    "market",
    "visa",
    "application",
    "admission",
    "requirements",
    "living",
    "accommodation",
    "housing",
    "expenses",
]

APPLICATION_CHECKLIST = """
# University Application Checklist

## Essential Documents
- [ ] Statement of Purpose (SOP)
- [ ] Letters of Recommendation (LOR)
- [ ] Official Transcripts
- [ ] Standardized Test Scores (GRE/GMAT if required)
- [ ] English Proficiency Test (TOEFL/IELTS)
- [ ] Resume/CV
- [ ] Financial Documents

## Additional Requirements
- [ ] Portfolio (if applicable)
- [ ] Writing Samples (if required)
- [ ] Research Proposal (for research programs)
- [ ] Letter of Intent

## Administrative Tasks
- [ ] Application Form Completion
- [ ] Application Fee Payment
- [ ] Passport Copy
- [ ] Photograph according to specifications
- [ ] Verification of Document Copies

## Post-Admission Steps
- [ ] Accept Admission Offer
- [ ] Pay Deposit
- [ ] Apply for Visa
- [ ] Book Accommodation
- [ ] Purchase Health Insurance
- [ ] Plan Travel
"""

# User data file path
USER_DATA_FILE = "user_data.json"

# Chatbot usage tips
CHATBOT_TIPS = """
### 💡 Tips for Using COMPASS

1. *Ask Specific Questions*
   - About university programs
   - About living costs in different locations
   - About job prospects in your field
   - About weather conditions

2. *Get Recommendations*
   - Click "Top 3 Recommendations" for personalized suggestions
   - Ask follow-up questions about specific universities
   - Inquire about admission requirements

3. *Explore Details*
   - Ask about specific universities
   - Request cost breakdowns
   - Learn about campus life
   - Get weather information

4. *Example Questions*
   - "Tell me more about Georgia Tech's program"
   - "What are the living costs in Atlanta?"
   - "How's the job market for data science in the Northeast?"
   - "Compare the weather between Boston and Miami"
"""


def load_user_data() -> dict:
    """Load user data from JSON file."""
    if os.path.exists(USER_DATA_FILE):
        try:
            with open(USER_DATA_FILE, "r") as f:
                return json.load(f)
        except Exception as e:
            logger.error(f"Error loading user data: {str(e)}")
    return {}


def save_user_data(data: dict):
    """Save user data to JSON file."""
    try:
        with open(USER_DATA_FILE, "w") as f:
            json.dump(data, f, indent=4)
    except Exception as e:
        logger.error(f"Error saving user data: {str(e)}")


def authenticate_user(username: str) -> bool:
    """Authenticate user and load their data."""
    users = load_user_data()

    # Initialize session state for new user
    if username not in users:
        users[username] = {
            "preferences": None,
            "chat_history": [],
            "last_recommendations": None,
            "applications": {},
            "created_at": datetime.now().isoformat(),
        }
        save_user_data(users)

    # Load user data into session state
    st.session_state.current_user = username
    st.session_state.user_data = users[username]
    st.session_state.authenticated = True

    return True


def save_user_preferences(username: str, preferences: dict):
    """Save user preferences."""
    users = load_user_data()
    if username in users:
        users[username]["preferences"] = preferences
        save_user_data(users)
        st.session_state.user_data = users[username]


def save_chat_history(username: str, chat_history: list):
    """Save chat history for a user."""
    users = load_user_data()
    if username in users:
        users[username]["chat_history"] = chat_history
        save_user_data(users)


def save_last_recommendations(username: str, recommendations: str):
    """Save last recommendations for context."""
    users = load_user_data()
    if username in users:
        users[username]["last_recommendations"] = recommendations
        save_user_data(users)
        st.session_state.user_data = users[username]


def update_application_tracker(username: str, university: str, status: str) -> None:
    """Update the user's university application tracking."""
    try:
        # Load current user data
        users = load_user_data()

        # Ensure the username exists
        if username not in users:
            users[username] = {
                "preferences": None,
                "chat_history": [],
                "last_recommendations": None,
                "applications": {},
                "created_at": datetime.now().isoformat(),
            }

        # Ensure applications dict exists
        if "applications" not in users[username]:
            users[username]["applications"] = {}

        # Update the application
        users[username]["applications"][university] = {
            "status": status,
            "updated_at": datetime.now().isoformat(),
        }

        # Save to file
        save_user_data(users)

        # Update session state
        st.session_state.user_data = users[username]

        # Log successful update
        logger.info(f"Updated application for {username}: {university} - {status}")

    except Exception as e:
        logger.error(f"Error updating application tracker: {str(e)}")
        raise e


def get_application_status(username: str) -> dict:
    """Get the user's university application status."""
    # First check session state for most up-to-date data
    if hasattr(st.session_state, "user_data"):
        return st.session_state.user_data.get("applications", {})

    # Fallback to file storage
    users = load_user_data()
    if username in users:
        return users[username].get("applications", {})
    return {}


def load_word_document(file_path: str) -> str:
    """Load content from a Word document."""
    try:
        doc = docx.Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        st.error(f"Error loading Word document: {str(e)}")
        return ""


def is_relevant_query(query: str) -> bool:
    """Check if the query is relevant to the system's purpose."""
    query_words = set(query.lower().split())
    return any(topic in query.lower() for topic in ALLOWED_TOPICS)


def generate_checklist_docx():
    """Generate and return a DOCX file for the application checklist."""
    try:
        doc = Document()
        doc.add_heading("University Application Checklist", 0)

        # Essential Documents Section
        doc.add_heading("Essential Documents", level=1)
        essential_items = [
            "Statement of Purpose (SOP)",
            "Letters of Recommendation (LOR)",
            "Official Transcripts",
            "Standardized Test Scores (GRE/GMAT if required)",
            "English Proficiency Test (TOEFL/IELTS)",
            "Resume/CV",
            "Financial Documents",
        ]
        for item in essential_items:
            doc.add_paragraph(item, style="List Bullet")

        # Additional Requirements Section
        doc.add_heading("Additional Requirements", level=1)
        additional_items = [
            "Portfolio (if applicable)",
            "Writing Samples (if required)",
            "Research Proposal (for research programs)",
            "Letter of Intent",
        ]
        for item in additional_items:
            doc.add_paragraph(item, style="List Bullet")

        # Administrative Tasks Section
        doc.add_heading("Administrative Tasks", level=1)
        admin_items = [
            "Application Form Completion",
            "Application Fee Payment",
            "Passport Copy",
            "Photograph according to specifications",
            "Verification of Document Copies",
        ]
        for item in admin_items:
            doc.add_paragraph(item, style="List Bullet")

        # Post-Admission Steps Section
        doc.add_heading("Post-Admission Steps", level=1)
        post_items = [
            "Accept Admission Offer",
            "Pay Deposit",
            "Apply for Visa",
            "Book Accommodation",
            "Purchase Health Insurance",
            "Plan Travel",
        ]
        for item in post_items:
            doc.add_paragraph(item, style="List Bullet")

        # Save to BytesIO object
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes
    except Exception as e:
        logger.error(f"Error generating checklist docx: {str(e)}")
        return None


def handle_tracker_command(query: str) -> str:
    """Handle natural language commands for the application tracker."""
    query = query.lower()
    if "add" in query and "to the tracker" in query:
        try:
            start_idx = query.find("add") + 4
            end_idx = query.find("to the tracker")
            if start_idx < end_idx:
                university = query[start_idx:end_idx].strip()
                update_application_tracker(
                    st.session_state.current_user, university, "Planning to Apply"
                )
                return f"Added {university} to your application tracker."
        except Exception as e:
            logger.error(f"Error processing tracker command: {str(e)}")
    return None


def initialize_chromadb():
    """Initialize ChromaDB with OpenAI embeddings."""
    global university_collection, living_expenses_collection, employment_collection
    global chroma_client, embedding_function

    try:
        # Setup OpenAI embedding function
        embedding_function = embedding_functions.OpenAIEmbeddingFunction(
            api_key=st.secrets["open-key"], model_name="text-embedding-ada-002"
        )

        # Connect to ChromaDB
        chroma_client = chromadb.PersistentClient(path="./chroma_db")

        # Check or create collections
        university_collection = chroma_client.get_or_create_collection(
            name="university_info", embedding_function=embedding_function
        )

        living_expenses_collection = chroma_client.get_or_create_collection(
            name="living_expenses", embedding_function=embedding_function
        )

        employment_collection = chroma_client.get_or_create_collection(
            name="employment_projections", embedding_function=embedding_function
        )

        logger.info("Successfully connected to or created collections in ChromaDB")
        return True

    except Exception as e:
        logger.error(f"Error initializing ChromaDB: {str(e)}")
        st.error(f"Failed to initialize database: {str(e)}")
        return False


def load_initial_data():
    """Load initial data into ChromaDB collections."""
    try:
        # Load datasets
        living_expenses_df = pd.read_csv(
            os.path.join("data", "Avg_Living_Expenses.csv")
        )
        employment_df = pd.read_csv(os.path.join("data", "Employment_Projections.csv"))
        university_text = load_word_document(
            os.path.join("data", "University_Data.docx")
        )

        # Process living expenses
        logger.info("Processing living expenses data...")
        for idx, row in living_expenses_df.iterrows():
            content = (
                f"State: {row['State']}\n"
                f"Cost of Living Index: {row['Index']}\n"
                f"Grocery: {row['Grocery']}\n"
                f"Housing: {row['Housing']}\n"
                f"Utilities: {row['Utilities']}\n"
                f"Transportation: {row['Transportation']}\n"
                f"Health: {row['Health']}\n"
                f"Miscellaneous: {row['Misc.']}"
            )
            living_expenses_collection.add(
                documents=[content],
                metadatas=[
                    {
                        "state": row["State"].strip(),
                        "type": "living_expenses",
                        "index": float(row["Index"]),
                    }
                ],
                ids=[f"living_expenses_{idx}"],
            )

        # Process employment projections
        logger.info("Processing employment projections data...")
        for idx, row in employment_df.iterrows():
            content = (
                f"Occupation: {row['Occupation Title']}\n"
                f"Employment 2023: {row['Employment 2023']}\n"
                f"Growth Rate: {row['Employment Percent Change, 2023-2033']}%\n"
                f"Annual Openings: {row['Occupational Openings, 2023-2033 Annual Average']}\n"
                f"Median Wage: ${row['Median Annual Wage 2023']}\n"
                f"Required Education: {row['Typical Entry-Level Education']}"
            )
            employment_collection.add(
                documents=[content],
                metadatas=[
                    {
                        "occupation": row["Occupation Title"],
                        "type": "employment",
                        "median_wage": float(row["Median Annual Wage 2023"]),
                    }
                ],
                ids=[f"employment_{idx}"],
            )

        # Process university data
        logger.info("Processing university data...")
        chunk_size = 1000
        chunks = [
            university_text[i : i + chunk_size]
            for i in range(0, len(university_text), chunk_size)
        ]

        for idx, chunk in enumerate(chunks):
            # Add some basic text preprocessing
            chunk = chunk.strip()
            if not chunk:  # Skip empty chunks
                continue

            university_collection.add(
                documents=[chunk],
                metadatas=[
                    {"chunk_id": idx, "type": "university", "length": len(chunk)}
                ],
                ids=[f"university_{idx}"],
            )

        logger.info("Successfully loaded initial data into ChromaDB")

    except Exception as e:
        logger.error(f"Error loading initial data: {str(e)}")
        raise e


def reset_chromadb():
    """Reset ChromaDB collections (useful for testing)."""
    global chroma_client
    try:
        if chroma_client:
            for collection_name in [
                "university_info",
                "living_expenses",
                "employment_projections",
            ]:
                try:
                    chroma_client.delete_collection(collection_name)
                    logger.info(f"Deleted collection: {collection_name}")
                except:
                    pass
        return initialize_chromadb()
    except Exception as e:
        logger.error(f"Error resetting ChromaDB: {str(e)}")
        return False


def get_living_expenses(state: str) -> str:
    """Query living expenses information."""
    try:
        results = living_expenses_collection.query(query_texts=[state], n_results=1)
        return (
            results["documents"][0][0]
            if results["documents"][0]
            else "No information found."
        )
    except Exception as e:
        logger.error(f"Error in get_living_expenses: {str(e)}")
        return f"Error retrieving living expenses: {str(e)}"


def get_job_market_trends(field: str) -> str:
    """Query job market trends."""
    try:
        results = employment_collection.query(query_texts=[field], n_results=3)
        return "\n\n".join(results["documents"][0])
    except Exception as e:
        logger.error(f"Error in get_job_market_trends: {str(e)}")
        return f"Error retrieving job market trends: {str(e)}"


def get_university_info(query: str) -> str:
    """Query university information."""
    try:
        results = university_collection.query(query_texts=[query], n_results=3)
        return "\n\n".join(results["documents"][0])
    except Exception as e:
        logger.error(f"Error in get_university_info: {str(e)}")
        return f"Error retrieving university information: {str(e)}"


def generate_application_tracker_template():
    """Generate a CSV template for tracking university applications."""
    template_df = pd.DataFrame(
        {
            "University Name": [
                "Example University",
                "Sample College",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
            ],
            "Program/Major": [
                "Computer Science",
                "Data Science",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
            ],
            "Application Status": [
                "In Progress",
                "Planning to Apply",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
            ],
            "Application Deadline": [
                "2024-12-31",
                "2024-11-30",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
            ],
            "Required Documents": [
                "Transcripts, LORs, SOP",
                "Transcripts, Portfolio",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
            ],
            "Notes": [
                "Need to request transcripts",
                "Portfolio in progress",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
            ],
            "Estimated Cost": ["$50,000", "$45,000", "", "", "", "", "", "", "", ""],
        }
    )

    # Convert to CSV bytes for download
    csv_bytes = template_df.to_csv(index=False).encode("utf-8")
    return csv_bytes


def get_weather_info(location: str) -> str:
    """Get weather information for a location's major city."""
    try:
        city = LOCATION_TO_CITY.get(location, location)
        base_url = "http://api.openweathermap.org/data/2.5/weather"
        params = {
            "q": f"{city},US",
            "appid": st.secrets["open-weather"],
            "units": "imperial",
        }

        response = requests.get(base_url, params=params)
        response.raise_for_status()

        data = response.json()
        return (
            f"Current weather in {city}: "
            f"Temperature: {data['main']['temp']}°F, "
            f"Feels like: {data['main']['feels_like']}°F, "
            f"Humidity: {data['main']['humidity']}%, "
            f"Conditions: {data['weather'][0]['description']}"
        )
    except Exception as e:
        logger.error(f"Error in get_weather_info: {str(e)}")
        return f"Error retrieving weather information: {str(e)}"


def get_top_recommendations() -> str:
    """Generate top 3 university recommendations based on user preferences."""
    try:
        preferences = st.session_state.user_data["preferences"]
        if not preferences:
            return (
                "Please set your preferences first to get personalized recommendations."
            )

        prompt = f"""Generate top 3 university recommendations based on these preferences:
        - Field of Study: {preferences['field_of_study']}
        - Budget Range: ${preferences['budget_min']}-${preferences['budget_max']}
        - Preferred Regions: {', '.join(preferences['preferred_locations'])}
        - Weather Preference: {preferences['weather_preference']}

        Format each recommendation as:
        [Number]. [University Name]
        * Key strengths in {preferences['field_of_study']}
        * Cost and aid highlights
        * Location and weather notes
        * Brief distinguishing features

        Keep each university description detailed but concise. Include specific programs, costs, and unique features. Make sure the text font is concistent enough, don't change the font mid sentence."""

        client = OpenAI(api_key=st.secrets["open-key"])
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are a university advisor providing detailed, specific recommendations.",
                },
                {"role": "user", "content": prompt},
            ],
            max_tokens=1000,
            temperature=0.7,
        )

        recommendations = response.choices[0].message.content
        save_last_recommendations(st.session_state.current_user, recommendations)
        return recommendations

    except Exception as e:
        logger.error(f"Error generating top recommendations: {str(e)}")
        return "Error generating recommendations. Please try again."


def get_recommendations(query: str) -> str:
    """Generate recommendations based on user query and preferences."""
    try:
        # Add handling for application tracker queries
        if "application tracker" in query.lower() or "my applications" in query.lower():
            applications = st.session_state.user_data.get("applications", {})
            if not applications:
                return "You haven't added any applications to track yet. You can add applications using the Application Tracker button."

            response = "Here are your current applications:\n\n"
            for univ, details in applications.items():
                response += f"• {univ}: {details['status']}\n"
            return response

        # Check if query is about a previously recommended university
        last_recommendations = st.session_state.user_data.get(
            "last_recommendations", ""
        )
        context = []

        if last_recommendations and any(
            university.lower() in query.lower()
            for university in ["georgia tech", "unc", "vanderbilt", "chapel hill"]
        ):
            context.append(("Previous Recommendations:", last_recommendations))

        # Get university information
        uni_info = get_university_info(query)
        if uni_info:
            context.append(("University Information:", uni_info))

        # Get living expenses if state is mentioned
        if any(
            word in query.lower() for word in ["state", "cost", "living", "expensive"]
        ):
            expenses_info = get_living_expenses(query)
            if expenses_info:
                context.append(("Living Expenses:", expenses_info))

        # Get job market trends if career/job is mentioned
        if any(
            word in query.lower()
            for word in ["job", "career", "employment", "salary", "work"]
        ):
            job_info = get_job_market_trends(
                st.session_state.user_data["preferences"]["field_of_study"]
            )
            if job_info:
                context.append(("Job Market Trends:", job_info))

        # Get weather for relevant locations
        if st.session_state.user_data["preferences"]["preferred_locations"]:
            location_weather = []
            for location in st.session_state.user_data["preferences"][
                "preferred_locations"
            ]:
                weather_info = get_weather_info(location)
                if not weather_info.startswith("Error"):
                    location_weather.append(weather_info)
            if location_weather:
                context.append(("Weather Information:", "\n".join(location_weather)))

        # Prepare prompt with context and user preferences
        context_text = "\n\n".join([f"{title}\n{info}" for title, info in context])
        preferences = st.session_state.user_data["preferences"]

        # Calculate appropriate token limit based on context length
        expected_length = len(context_text) + len(query)
        max_tokens = 500 if expected_length < 1000 else 300

        prompt = f"""As a university advisor, help with this query: {query}

Context:
{context_text}

Student Profile:
- Field: {preferences['field_of_study']}
- Budget: ${preferences['budget_min']}-${preferences['budget_max']}
- Locations: {', '.join(preferences['preferred_locations'])}
- Weather: {preferences['weather_preference']}

Previous conversation context (if relevant):
{last_recommendations}

Provide a concise, specific response focusing on the query. If the query is about a specific university from the previous recommendations, reference that information. If the query is completely unrelated alert the user saying you can only answer questions on universities, job aspects"""

        client = OpenAI(api_key=st.secrets["open-key"])
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a helpful university advisor for international students. "
                        "Provide specific, actionable insights. If the response needs to be "
                        "longer than the token limit allows, provide a concise summary of "
                        "the most important points."
                    ),
                },
                {"role": "user", "content": prompt},
            ],
            max_tokens=max_tokens,
            temperature=0.7,
        )

        return response.choices[0].message.content

    except Exception as e:
        logger.error(f"Error in get_recommendations: {str(e)}")
        return "I apologize, but I encountered an error. Please try asking a more specific question."


def initialize_session_state():
    """Initialize all session state variables."""
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    if "current_user" not in st.session_state:
        st.session_state.current_user = None
    if "user_data" not in st.session_state:
        st.session_state.user_data = None
    if "initialized" not in st.session_state:
        st.session_state.initialized = False
    if "show_preferences" not in st.session_state:
        st.session_state.show_preferences = False


def show_login_page():
    """Display a login page with logo."""
    # Custom CSS to ensure dark theme
    st.markdown(
        """
        <style>
        /* Dark theme overrides */
        .stTextInput input {
            background-color: #2b303b;
            color: white;
            border: 1px solid #4a4d52;
        }
        .stTextInput input::placeholder {
            color: #8e929b;
        }
        .stButton button {
            width: 100%;
            background-color: #ff4b4b;
            color: white;
            border: none;
            padding: 10px 20px;
            border-radius: 5px;
            transition: all 0.3s ease;
        }
        .stButton button:hover {
            background-color: #ff3333;
            border: none;
        }
        div[data-testid="stToolbar"] {
            display: none;
        }
        .main > div {
            background-color: #1b1e23;
        }
        /* Image container styling */
        .img-container {
            display: flex;
            justify-content: center;
            margin-bottom: 1rem;
        }
        </style>
    """,
        unsafe_allow_html=True,
    )

    col1, col2, col3 = st.columns([1, 2, 1])

    with col2:
        # Image display
        try:
            st.image(
                os.path.join("data", "compass_logo_wide.png"), use_container_width=True
            )
        except Exception as e:
            st.error(f"Unable to load image: {str(e)}")

        st.markdown(
            """
            <div style='text-align: center; padding: 20px; color: white;'>
                <h1>🎓 Welcome to COMPASS</h1>
                <p style='font-size: 1.2em; color: #8e929b;'>Your AI-powered University Guide</p>
            </div>
        """,
            unsafe_allow_html=True,
        )

        with st.form("login_form", clear_on_submit=True):
            username = st.text_input(
                "👤 Enter your name",
                placeholder="Type your name here...",
                help="This will be used as your username & will personalize your experience",
            )

            submitted = st.form_submit_button("🚀 Start Your Journey")

            if submitted and username:
                if authenticate_user(username):
                    st.success(f"Welcome aboard, {username}! 🎉")
                    if not st.session_state.user_data.get("preferences"):
                        st.info("Let's set up your preferences to get started.")
                    st.rerun()

        st.markdown(
            """
            <div style='text-align: center; padding: 20px; color: #8e929b;'>
                <p>COMPASS helps you:</p>
                <p style='color: #b4b9c2;'>🎯 Find the perfect university</p>
                <p style='color: #b4b9c2;'>💰 Understand costs and living expenses</p>
                <p style='color: #b4b9c2;'>🌤 Get location and weather insights</p>
                <p style='color: #b4b9c2;'>💼 Explore career opportunities</p>
            </div>
        """,
            unsafe_allow_html=True,
        )


def show_preferences_form(existing_preferences=None):
    # Initialize session state variables for toggling widgets
    if "use_all_locations" not in st.session_state:
        st.session_state.use_all_locations = False
    if "use_all_weather" not in st.session_state:
        st.session_state.use_all_weather = False

    # Field of Study
    field_of_study = st.text_input(
        "Field of Study",
        value=existing_preferences.get("field_of_study", "") if existing_preferences else "",
        help="Enter your intended field of study",
    )

    # Budget Range
    default_budget_any = (
        existing_preferences.get("budget_min", 0) == 0 and
        existing_preferences.get("budget_max", 100000) == 100000
        if existing_preferences else False
    )
    use_any_budget = st.checkbox(
        "Any Budget",
        value=default_budget_any,
        help="Select this to consider all budget ranges",
    )
    if not use_any_budget:
        budget_range = st.slider(
            "Budget Range (USD/Year)",
            0,
            100000,
            (
                existing_preferences.get("budget_min", 20000),
                existing_preferences.get("budget_max", 50000),
            ) if existing_preferences else (20000, 50000),
            help="Select your annual budget range",
        )
    else:
        budget_range = (0, 100000)
        st.info("Considering all budget ranges")

    # Locations
    st.session_state.use_all_locations = st.checkbox(
        "All Regions",
        value=st.session_state.use_all_locations,
        help="Select this to consider all regions",
    )
    if not st.session_state.use_all_locations:
        preferred_locations = st.multiselect(
            "Preferred Regions",
            ["Northeast", "Southeast", "Midwest", "Southwest", "West Coast"],
            default=existing_preferences.get("preferred_locations", []) if existing_preferences else [],
            help="Select your preferred regions",
        )
    else:
        preferred_locations = ["Northeast", "Southeast", "Midwest", "Southwest", "West Coast"]
        st.info("Considering all regions")

    # Weather Preferences
    st.session_state.use_all_weather = st.checkbox(
        "All Weather Types",
        value=st.session_state.use_all_weather,
        help="Select this to consider all weather types",
    )
    if not st.session_state.use_all_weather:
        weather_preferences = st.multiselect(
            "Weather Preferences",
            ["Cold", "Warm", "Hot"],
            default=existing_preferences.get("weather_preference", []) if existing_preferences else [],
            help="Select multiple weather preferences",
        )
    else:
        weather_preferences = ["Cold", "Warm", "Hot"]
        st.info("Considering all weather types")

    # Save Preferences Button
    if st.button("Save Preferences"):
        # Validate and save preferences
        preferences = {
            "field_of_study": field_of_study,
            "budget_min": budget_range[0],
            "budget_max": budget_range[1],
            "preferred_locations": preferred_locations,
            "weather_preference": weather_preferences,
        }
        save_user_preferences(st.session_state.current_user, preferences)
        st.success("Preferences saved successfully!")
        st.success("✅ Preferences saved successfully!")
            
            st.session_state.show_preferences = False
            st.rerun()
            
            return True
    return False


def show_application_tracker():
    """Display the application tracking interface."""
    st.subheader("📝 Application Tracker Template")

    # Generate and provide download button for CSV template
    csv_template = generate_application_tracker_template()
    st.download_button(
        label="📥 Download Application Tracker Template",
        data=csv_template,
        file_name="university_application_tracker.csv",
        mime="text/csv",
        help="Download a CSV template to track your university applications",
    )

    st.write(
        """
    ### How to Use the Template:
    1. Download the CSV template
    2. Open it in Excel, Google Sheets, or any spreadsheet software
    3. Fill in your application details
    4. Keep track of deadlines and required documents
    5. Update status as you progress
    
    The template includes example entries to help you get started.
    """
    )


def authenticate_user(username: str) -> bool:
    """Authenticate user and load their data."""
    users = load_user_data()

    # Initialize session state for new user with explicit applications field
    if username not in users:
        users[username] = {
            "preferences": None,
            "chat_history": [],
            "last_recommendations": None,
            "applications": {},  # Explicitly initialize applications
            "created_at": datetime.now().isoformat(),
        }
        save_user_data(users)

    # Load user data into session state
    st.session_state.current_user = username
    st.session_state.user_data = users[username]
    st.session_state.authenticated = True

    return True


def show_sidebar():
    """Display sidebar with tips and preferences."""
    with st.sidebar:
        try:
            st.image(os.path.join("data", "compass_logo.png"), use_container_width=True)
        except:
            st.error("Image not found. Please check the path.")

        # Show current user and date
        current_date = datetime.now().strftime("%B %d, %Y")
        st.write(f"📅 {current_date}")
        st.write(f"👤 Logged in as: {st.session_state.current_user}")

        # Separator
        st.markdown("---")

        # Current Preferences Section
        st.header("📋 Your Preferences")
        prefs = st.session_state.user_data.get("preferences", {})
        if prefs:
            st.write(f"*Field of Study:* {prefs.get('field_of_study')}")

            # Display budget
            if prefs.get("budget_min") == 0 and prefs.get("budget_max") == 100000:
                st.write("*Budget:* Any")
            else:
                st.write(
                    f"*Budget:* ${prefs.get('budget_min'):,} - ${prefs.get('budget_max'):,}"
                )

            # Display locations
            locations = prefs.get("preferred_locations", [])
            if len(locations) == 5:
                st.write("*Locations:* All")
            else:
                st.write(f"*Locations:* {', '.join(locations)}")

            # Display weather preferences
            weather_pref = prefs.get("weather_preference")
            if weather_pref == "All" or (
                isinstance(weather_pref, list) and len(weather_pref) == 4
            ):
                st.write("*Weather:* All")
            elif isinstance(weather_pref, list):
                st.write(f"*Weather:* {', '.join(weather_pref)}")
            else:
                st.write("*Weather:* None")

            if st.button("✏ Edit Preferences"):
                st.session_state.show_preferences = True

        # Add introduction
        st.markdown(
            """
        ---
        ### About COMPASS
        Your AI companion for university selection and application guidance. Get personalized recommendations for universities, and receive insights about living costs, weather, and job prospects in different locations.
        """
        )
        # Tips Section
        st.markdown(CHATBOT_TIPS)

        # Logout button at the bottom
        if st.button("🚪 Logout", key="logout_sidebar"):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()


def show_chat_interface():
    """Display the main chat interface."""
    # Show preferences edit form if requested
    if st.session_state.show_preferences:
        st.header("✏ Edit Preferences")
        if show_preferences_form(st.session_state.user_data.get("preferences")):
            st.session_state.show_preferences = False
            st.rerun()
        return

    st.header("💬 Chat with COMPASS")

    # Buttons row
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        if st.button("🌟 Top 3 Recommendations"):
            recommendations = get_top_recommendations()
            st.session_state.user_data["chat_history"].append(
                {"role": "assistant", "content": recommendations}
            )
            save_chat_history(
                st.session_state.current_user,
                st.session_state.user_data["chat_history"],
            )

    with col2:
        checklist_doc = generate_checklist_docx()
        if checklist_doc:
            st.download_button(
                label="📋 Download Checklist",
                data=checklist_doc,
                file_name="application_checklist.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    with col3:
        if st.button("📊 Tracker Template"):
            show_application_tracker()

    with col4:
        if st.button("🗑 Clear Chat"):
            st.session_state.user_data["chat_history"] = []
            save_chat_history(st.session_state.current_user, [])
            st.rerun()

    # Display chat history
    for message in st.session_state.user_data["chat_history"]:
        with st.chat_message(message["role"]):
            st.write(message["content"])

    # Chat input
    if prompt := st.chat_input(
        "Ask me about universities, programs, costs, or job prospects..."
    ):
        # Check for tracker commands
        tracker_response = handle_tracker_command(prompt)
        if tracker_response:
            with st.chat_message("assistant"):
                st.write(tracker_response)
            st.session_state.user_data["chat_history"].append(
                {"role": "assistant", "content": tracker_response}
            )
            save_chat_history(
                st.session_state.current_user,
                st.session_state.user_data["chat_history"],
            )
            st.rerun()
            return

        # Display user message
        with st.chat_message("user"):
            st.write(prompt)

        # Add to chat history
        st.session_state.user_data["chat_history"].append(
            {"role": "user", "content": prompt}
        )

        # Get and display assistant response
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                response = get_recommendations(prompt)
                st.write(response)

                # Add to chat history
                st.session_state.user_data["chat_history"].append(
                    {"role": "assistant", "content": response}
                )

        # Save updated chat history
        save_chat_history(
            st.session_state.current_user, st.session_state.user_data["chat_history"]
        )


def main():
    """Main application function."""
    # Page configuration
    st.set_page_config(
        page_title="COMPASS - University Recommendation System",
        page_icon="🎓",
        layout="wide",
        initial_sidebar_state="expanded",
    )

    # Initialize session state
    initialize_session_state()

    # Initialize ChromaDB if not already initialized
    if not st.session_state.initialized:
        with st.spinner("Initializing system..."):
            if initialize_chromadb():
                st.session_state.initialized = True
            else:
                st.error("Failed to initialize the system. Please refresh the page.")
                return

    # Display appropriate page based on authentication state
    if not st.session_state.authenticated:
        show_login_page()
    else:
        # Show sidebar
        show_sidebar()

        # Show preferences page if preferences not set
        if not st.session_state.user_data.get("preferences"):
            show_preferences_form()
        else:
            show_chat_interface()

if __name__ == "__main__":
    main()